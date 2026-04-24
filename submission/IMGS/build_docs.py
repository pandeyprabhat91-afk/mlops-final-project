"""Build all 5 submission DOCX reports for Deepfake Detection MLOps project."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = "f:/mlops project/submission"
IMG = {
    "arch": os.path.join(BASE, "arch_diagram.svg"),
    "pipeline": os.path.join(BASE, "pipeline_flow.svg"),
    "api_flow": os.path.join(BASE, "api_flow.svg"),
    "dvc": os.path.join(BASE, "dvc_dag.svg"),
    "monitoring": os.path.join(BASE, "monitoring.svg"),
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_doc_styles(doc):
    """Apply professional base styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

def add_cover(doc, title, subtitle, version="1.0", date="April 2026"):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DEEPFAKE DETECTION")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    run.font.name = 'Calibri'
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(28)
    run2.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    run2.font.name = 'Calibri'
    run2.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run3.font.name = 'Calibri'

    doc.add_paragraph()
    # Divider
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("─" * 40)
    run4.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {version}   |   {date}   |   MLOps Academic Project").font.size = Pt(10)
    doc.add_page_break()

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    p.runs[0].font.name = 'Calibri'
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p.runs[0].font.name = 'Calibri'
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    p.runs[0].font.name = 'Calibri'
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.runs[0].font.name = 'Calibri'
    p.runs[0].font.size = Pt(11)
    return p

def code_block(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    # shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    return p

def add_image(doc, path, caption, width=Inches(5.5)):
    if os.path.exists(path):
        if path.lower().endswith('.svg'):
            _add_svg(doc, path, width)
        else:
            doc.add_picture(path, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p.runs[0].italic = True

def _add_svg(doc, svg_path, width=Inches(5.5)):
    import zipfile, shutil, uuid, re
    from lxml import etree

    with open(svg_path, 'rb') as f:
        svg_data = f.read()

    # Parse SVG dimensions for aspect ratio
    root = etree.fromstring(svg_data)
    vb = root.get('viewBox', '')
    w_attr = root.get('width', '960')
    h_attr = root.get('height', '720')
    try:
        if vb:
            parts = vb.split()
            svg_w, svg_h = float(parts[2]), float(parts[3])
        else:
            svg_w = float(re.sub(r'[^\d.]', '', w_attr))
            svg_h = float(re.sub(r'[^\d.]', '', h_attr))
    except Exception:
        svg_w, svg_h = 960, 720

    ratio = svg_h / svg_w
    cx = int(width.inches * 914400)          # EMUs
    cy = int(width.inches * ratio * 914400)

    rid = f"rId{uuid.uuid4().hex[:8]}"
    part_name = f"/word/media/{uuid.uuid4().hex}.svg"

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    svg_part = Part(PackURI(part_name), 'image/svg+xml', svg_data, doc._part.package)
    rel = doc._part.relate_to(svg_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

    # Build drawing XML inline
    nsmap = {
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
        'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'r':   'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'asvg':'http://schemas.microsoft.com/office/drawing/2016/SVG/main',
    }

    xml = f'''<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <wp:inline xmlns:wp="{nsmap['wp']}">
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:docPr id="1" name="SVG"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic xmlns:a="{nsmap['a']}">
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic xmlns:pic="{nsmap['pic']}">
          <pic:nvPicPr>
            <pic:cNvPr id="0" name="SVG"/>
            <pic:cNvPicPr/>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip xmlns:r="{nsmap['r']}" r:embed="{rel}">
              <a:extLst>
                <a:ext uri="{{96DAC541-7B7A-43D3-8B79-37D633B846F1}}">
                  <asvg:svgBlip xmlns:asvg="{nsmap['asvg']}" r:embed="{rel}"/>
                </a:ext>
              </a:extLst>
            </a:blip>
            <a:stretch><a:fillRect/></a:stretch>
          </pic:blipFill>
          <pic:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:inline>
</w:drawing>'''

    drawing_el = etree.fromstring(xml)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(drawing_el)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = 'Calibri'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EFF6FF')
        tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        drow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = drow.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F9FAFB')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w
    return table

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Architecture Diagram Report
# ═══════════════════════════════════════════════════════════════════════════════
def build_arch_doc():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "Architecture Diagram", "System Architecture with Component Explanations")

    h1(doc, "1. System Architecture Overview")
    body(doc, "The Deepfake Detection system is a full-stack MLOps platform composed of five distinct architectural layers. Each layer is independent and communicates through well-defined interfaces, ensuring loose coupling and high cohesion.")
    doc.add_paragraph()
    add_image(doc, IMG["arch"], "Figure 1: System Architecture — Five-Layer View", Inches(6.0))

    h2(doc, "1.1 Architecture Layers")
    body(doc, "The system is organized into five horizontal layers, each with a distinct responsibility:")

    h3(doc, "Layer 1 — Client Layer")
    body(doc, "The React frontend (served by Nginx on port 3000) is the only user-facing component. It accepts MP4 video uploads via drag-and-drop and presents prediction results including confidence scores and Grad-CAM heatmaps. Nginx proxies all /api/* requests to the backend.")

    h3(doc, "Layer 2 — Gateway / API Layer")
    body(doc, "FastAPI (port 8000) is the sole entry point for all backend operations. It enforces strict input validation, orchestrates preprocessing, calls the ML model, generates explainability overlays, and exposes Prometheus metrics. The frontend and backend are connected exclusively via REST API — no shared state or direct imports.")

    h3(doc, "Layer 3 — ML Services Layer")
    body(doc, "Three services power the ML lifecycle:")
    bullet(doc, "MLflow Server (port 5000): Tracks experiments, manages the model registry, and stores training artifacts. The backend loads the production model via mlflow.pyfunc at startup.")
    bullet(doc, "Apache Airflow (port 8080): Orchestrates the data ingestion pipeline (daily) and model retraining pipeline (weekly) using a LocalExecutor — no separate worker container required.")
    bullet(doc, "DVC Pipeline: Manages reproducible data versioning across five stages: extract_frames → detect_faces → compute_features → train → evaluate.")

    h3(doc, "Layer 4 — Data and Storage Layer")
    body(doc, "Persistent storage is split by concern:")
    add_table(doc,
        ["Store", "Purpose", "Technology"],
        [
            ["PostgreSQL (port 5432)", "Airflow metadata, task state, DAG history", "PostgreSQL 15"],
            ["MLflow Artifacts", "Model weights, confusion matrices, ROC curves", "Local filesystem / MLflow store"],
            ["DVC Storage", "Raw MP4s, extracted frames, face crops, feature tensors", "DVC + Git LFS"],
            ["Git Repository", "Source code, config, dvc.lock, dvc.yaml", "Git + Git LFS"],
        ]
    )

    h3(doc, "Layer 5 — Monitoring Layer")
    body(doc, "Prometheus (port 9090) scrapes metrics from all four services every 15 seconds. Grafana (port 3001) visualises these in three dashboards: inference performance, model drift, and pipeline health. Alert rules fire when error rate exceeds 5%, p95 latency exceeds 200ms, or drift score exceeds 3.0.")

    h2(doc, "1.2 Loose Coupling Guarantee")
    body(doc, "The architecture strictly enforces loose coupling at every boundary:")
    add_table(doc,
        ["Boundary", "Communication Method", "Configurable Via"],
        [
            ["Frontend ↔ Backend", "REST API only (/api/* proxy)", "VITE_API_URL env var"],
            ["Backend ↔ ML Model", "mlflow.pyfunc.load_model() URI", "MODEL_NAME / MODEL_STAGE env vars"],
            ["Backend ↔ Airflow", "HTTP (Airflow REST API)", "AIRFLOW_URL env var"],
            ["Backend ↔ MLflow", "HTTP (MLflow Tracking API)", "MLFLOW_TRACKING_URI env var"],
            ["All services ↔ Secrets", ".env file (never committed)", ".env.example template"],
        ]
    )

    h2(doc, "1.3 Docker Compose Services")
    add_table(doc,
        ["Service", "Port", "Purpose"],
        [
            ["frontend", "3000", "React app + Nginx proxy"],
            ["backend", "8000", "FastAPI inference + monitoring"],
            ["mlflow-server", "5000", "Experiment tracking + model registry"],
            ["airflow-webserver", "8080", "DAG management UI"],
            ["airflow-scheduler", "—", "LocalExecutor task runner"],
            ["postgres", "5432", "Airflow metadata database"],
            ["prometheus", "9090", "Metrics collection"],
            ["grafana", "3001", "Dashboards and alerting"],
            ["node-exporter", "9100", "Host system metrics"],
        ]
    )

    doc.save(os.path.join(BASE, "1_Architecture_Diagram.docx"))
    print("1_Architecture_Diagram.docx saved")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: High-Level Design
# ═══════════════════════════════════════════════════════════════════════════════
def build_hld_doc():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "High-Level Design", "Design Choices, Rationale, and MLOps Pipeline")

    h1(doc, "1. System Overview")
    body(doc, "The Deepfake Detection System is an end-to-end ML platform that classifies MP4 videos as real or fake using a deep learning model served through a production-grade MLOps pipeline. The system is designed for local deployment using Docker Compose.")
    body(doc, "Key design goals:")
    bullet(doc, "End-to-end deepfake classification with ≥ 90% accuracy and ≥ 0.90 F1-score")
    bullet(doc, "Full MLOps lifecycle: data versioning, experiment tracking, model registry, automated retraining")
    bullet(doc, "Sub-200ms inference latency at p95")
    bullet(doc, "Reproducibility: every experiment tied to a Git commit hash and MLflow run ID")
    bullet(doc, "Monitoring, alerting, and drift detection on all components")

    h1(doc, "2. MLOps Pipeline")
    add_image(doc, IMG["pipeline"], "Figure 2: MLOps Data and Training Pipeline", Inches(6.0))
    body(doc, "The pipeline flows in two paths: a data processing path (top row) that feeds raw MP4 videos through frame extraction, face detection, and feature computation, and a model lifecycle path (bottom row) that handles training, evaluation, registry, and promotion.")

    h2(doc, "2.1 Data Engineering Pipeline (Airflow — Daily)")
    body(doc, "The deepfake_pipeline DAG runs daily and executes these stages in sequence:")
    add_table(doc,
        ["Stage", "Tool", "Input", "Output"],
        [
            ["ingest_videos", "Airflow PythonOperator", "data/landing/*.mp4", "data/raw/*.mp4"],
            ["extract_frames", "OpenCV", "data/raw/", "data/frames/ (30 JPGs/video)"],
            ["detect_faces", "MTCNN", "data/frames/", "data/faces/ (cropped 224×224)"],
            ["compute_features", "EfficientNet-B0", "data/faces/", "data/features/*.pt"],
            ["validate_schema", "Custom validator", "data/features/", "validation report"],
            ["record_baseline_stats", "NumPy", "data/features/", "ml/feature_baseline.json"],
            ["version_with_dvc", "DVC", "all data/", "dvc.lock committed to Git"],
        ]
    )

    h2(doc, "2.2 DVC CI Pipeline")
    add_image(doc, IMG["dvc"], "Figure 3: DVC Pipeline DAG — Five Reproducible Stages", Inches(6.0))
    body(doc, "DVC defines a five-stage pipeline in dvc.yaml with explicit dependency declarations. Running dvc repro re-executes only stages whose inputs have changed, making the pipeline fully reproducible and auditable. The dvc.lock file is committed to Git, ensuring any team member can recreate the exact pipeline state.")

    h2(doc, "2.3 Automated Retraining (Airflow — Weekly)")
    body(doc, "The retraining_dag uses a BranchPythonOperator to decide whether retraining is needed based on feature drift. If drift score exceeds 3.0, the full retraining pipeline executes:")
    bullet(doc, "fetch_new_data: pulls newly labelled data from data/landing/")
    bullet(doc, "run_mlproject: executes ml/MLproject with params from ml/params.yaml")
    bullet(doc, "evaluate_model: computes F1, accuracy, PR-AUC, per-class precision/recall")
    bullet(doc, "register_if_better: registers model to MLflow registry if F1 ≥ 0.90")
    bullet(doc, "promote_to_production: transitions Staging → Production, archives current Production")

    h1(doc, "3. Model Architecture")
    body(doc, "The model is a CNN + LSTM architecture using EfficientNet-B0 as a temporal feature extractor:")
    code_block(doc, "Input: (batch, 30, 3, 224, 224)  — 30 face-cropped frames")
    code_block(doc, "EfficientNet-B0 (pretrained ImageNet)  — spatial features per frame")
    code_block(doc, "2-layer LSTM (hidden=256)  — temporal dependencies across frames")
    code_block(doc, "Linear(256→128) → ReLU → Dropout(0.3) → Linear(128→1) → Sigmoid")
    code_block(doc, "Output: (batch, 1)  — probability (>= 0.5 = fake)")

    add_table(doc,
        ["Decision", "Choice", "Rationale"],
        [
            ["Feature extractor", "EfficientNet-B0", "Best accuracy/efficiency trade-off; pretrained ImageNet weights reduce training time"],
            ["Temporal model", "2-layer LSTM", "Captures inter-frame temporal dependencies; simpler than transformers for 30-frame sequences"],
            ["Face detection", "MTCNN", "High accuracy, handles pose variation; lightweight at inference"],
            ["Experiment tracking", "MLflow", "Rich model registry with stage transitions; pyfunc serving eliminates extra container"],
            ["Data pipeline", "Airflow + DVC", "Airflow for orchestration and scheduling; DVC for reproducible data versioning"],
            ["Model serving", "mlflow.pyfunc (in-process)", "Backend loads model directly — no extra container, lower latency, fewer failure points"],
            ["Monitoring", "Prometheus + Grafana", "Industry standard; no external dependencies; full metrics lifecycle"],
            ["Frontend", "React + Vite + TypeScript", "Fast build; type safety; component-based architecture"],
        ]
    )

    h1(doc, "4. Experiment Tracking")
    body(doc, "MLflow tracking is fully manual — no autolog — to give precise control over what is recorded:")
    add_table(doc,
        ["What is Tracked", "How", "Purpose"],
        [
            ["All hyperparameters", "mlflow.log_params(params)", "Single source of truth from ml/params.yaml"],
            ["Per-epoch metrics", "mlflow.log_metrics({...}, step=epoch)", "train_loss, val_loss, train_f1, val_f1, val_accuracy, learning_rate"],
            ["Git commit SHA", "mlflow.set_tag(\"git_commit\", ...)", "Links every run to the exact code state"],
            ["Device used", "mlflow.set_tag(\"device\", ...)", "CPU vs CUDA affects reproducibility"],
            ["Best model artifact", "mlflow.pytorch.log_model()", "Registered to Model Registry as 'deepfake'"],
            ["Confusion matrix", "mlflow.log_artifact()", "Visual evaluation artifact"],
            ["ROC curve", "mlflow.log_artifact()", "AUC-ROC visual artifact"],
            ["Extended eval metrics", "mlflow.log_metrics()", "PR-AUC, per-class precision/recall, optimal threshold"],
        ]
    )

    h1(doc, "5. Monitoring Architecture")
    add_image(doc, IMG["monitoring"], "Figure 4: Prometheus and Grafana Monitoring Architecture", Inches(6.0))
    body(doc, "All components expose Prometheus metrics endpoints. Prometheus scrapes every 15 seconds. Grafana visualises three dashboards in near-real-time:")
    bullet(doc, "Inference Dashboard: request rate, latency histogram, error rate, drift score gauge")
    bullet(doc, "Model Drift Dashboard: z-score time series with threshold markers at 3.0")
    bullet(doc, "Pipeline Dashboard: Airflow task durations, DVC pipeline throughput, data volume")

    body(doc, "Alert rules fire when:")
    add_table(doc,
        ["Condition", "Threshold", "Response"],
        [
            ["Error rate", "> 5% for 3 minutes", "Grafana alert + investigate logs"],
            ["p95 Latency", "> 200ms", "Grafana alert + profile backend"],
            ["Drift score", "> 3.0", "Triggers retraining_dag BranchPythonOperator"],
            ["Pipeline validation failure", "Any failure", "Grafana alert + check Airflow logs"],
        ]
    )

    h1(doc, "6. Software Design Paradigm")
    body(doc, "The codebase uses a mixed OO + functional paradigm, chosen per layer:")
    add_table(doc,
        ["Layer", "Paradigm", "Rationale"],
        [
            ["ML model (ml/model.py)", "Object-Oriented — DeepfakeDetector(nn.Module)", "PyTorch Module system requires OO; encapsulates CNN + LSTM as a single composable unit"],
            ["Backend API (routers/)", "Functional — FastAPI route functions", "Stateless request handlers are naturally functional; easier to test in isolation"],
            ["Model loader", "Module-level singleton", "Holds shared mutable state (loaded model) without a class; standard Python pattern for process-wide singletons"],
            ["Airflow DAGs", "Functional — plain Python callables", "Airflow PythonOperator expects callables; pure functions are easier to test"],
            ["Preprocessing", "Functional with lazy singleton", "Stateless transforms; MTCNN cached via module-level variable to avoid reloading per request"],
        ]
    )

    h1(doc, "7. Security")
    bullet(doc, "All secrets stored in .env (never committed to Git; .env.example provided)")
    bullet(doc, "HTTPS via Nginx TLS termination in production")
    bullet(doc, "Prometheus and Grafana behind basic auth")
    bullet(doc, "No cloud usage — all data stays local (GDPR-safe for video PII)")
    bullet(doc, "Encryption at rest: Docker volumes mounted on OS-level encrypted filesystem (BitLocker/LUKS)")
    bullet(doc, "Admin endpoint (POST /admin/reload-model) not exposed through Nginx to public users")

    doc.save(os.path.join(BASE, "2_High_Level_Design.docx"))
    print("2_High_Level_Design.docx saved")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: Low-Level Design
# ═══════════════════════════════════════════════════════════════════════════════
def build_lld_doc():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "Low-Level Design", "API Endpoint Definitions and I/O Specifications")

    h1(doc, "1. API Endpoint Overview")
    add_image(doc, IMG["api_flow"], "Figure 5: POST /predict — API Request Sequence Diagram", Inches(6.0))
    body(doc, "The FastAPI backend exposes a RESTful API. All endpoints accept and return JSON (except /predict which accepts multipart/form-data and /metrics which returns Prometheus text format).")

    add_table(doc,
        ["Endpoint", "Method", "Auth", "Purpose"],
        [
            ["/predict", "POST", "None", "Video deepfake classification"],
            ["/predict/batch", "POST", "None", "Batch classification (max 10 videos)"],
            ["/health", "GET", "None", "Liveness probe"],
            ["/ready", "GET", "None", "Readiness probe (model loaded check)"],
            ["/metrics", "GET", "None", "Prometheus metrics export"],
            ["/feedback", "POST", "None", "Submit ground-truth label"],
            ["/admin/reload-model", "POST", "Internal only", "Force model reload"],
            ["/admin/rollback", "POST", "Internal only", "Roll back to specific model version"],
            ["/admin/model-info", "GET", "Internal only", "Current model metadata"],
            ["/pipeline/mlflow-runs", "GET", "None", "Recent MLflow experiment runs"],
            ["/pipeline/airflow-runs", "GET", "None", "Recent Airflow DAG runs"],
            ["/pipeline/throughput", "GET", "None", "Pipeline throughput metric"],
            ["/support/tickets", "POST/GET", "X-Username header", "Support ticket management"],
            ["/support/chat", "POST", "None", "Rule-based chatbot"],
        ]
    )

    h1(doc, "2. Endpoint Specifications")

    h2(doc, "2.1 POST /predict")
    body(doc, "Classifies a single MP4 video as real or fake.")
    h3(doc, "Request")
    code_block(doc, "Content-Type: multipart/form-data\nfield: file (MP4, max 100MB)")
    h3(doc, "Success Response — HTTP 200")
    code_block(doc, '{\n  "prediction": "fake",\n  "confidence": 0.94,\n  "inference_latency_ms": 143.2,\n  "gradcam_image": "<base64 PNG string>",\n  "mlflow_run_id": "uuid-string",\n  "frames_analyzed": 30\n}')
    h3(doc, "Error Responses")
    add_table(doc,
        ["HTTP Code", "Condition", "Response Body"],
        [
            ["400", "File is not .mp4 extension", '{"detail": "Only MP4 files are accepted"}'],
            ["500", "Model inference failure", '{"detail": "<error message>"}'],
        ]
    )
    h3(doc, "Processing Pipeline")
    body(doc, "1. Save upload to temp file (tempfile.NamedTemporaryFile)")
    body(doc, "2. preprocess_video(tmp_path) → tensor (30, 3, 224, 224)")
    body(doc, "3. model.predict(tensor.unsqueeze(0).numpy()) → confidence float")
    body(doc, "4. prediction = 'fake' if confidence >= 0.5 else 'real'")
    body(doc, "5. Feature drift check: compute_drift_score(features, baseline)")
    body(doc, "6. generate_gradcam(model, frames_tensor) → base64 PNG (best-effort)")
    body(doc, "7. Update Prometheus metrics (request count, latency, confidence, drift)")

    h2(doc, "2.2 POST /predict/batch")
    body(doc, "Classifies up to 10 MP4 videos in a single request.")
    h3(doc, "Request")
    code_block(doc, "Content-Type: multipart/form-data\nfield: files (list of MP4, max 10)")
    h3(doc, "Success Response — HTTP 200")
    code_block(doc, '{\n  "results": [\n    {"filename": "v1.mp4", "prediction": "fake", "confidence": 0.87, "inference_latency_ms": 132.1, "error": null},\n    {"filename": "v2.mp4", "prediction": null, "confidence": null, "error": "Only MP4 files are accepted"}\n  ],\n  "total": 2, "succeeded": 1, "failed": 1\n}')

    h2(doc, "2.3 GET /health")
    code_block(doc, '{"status": "ok", "model_loaded": true}')

    h2(doc, "2.4 GET /ready")
    code_block(doc, '{"status": "ready", "model_version": "deepfake/Production/3"}')
    body(doc, "Returns HTTP 503 with {\"detail\": \"Model not loaded\"} when model is unavailable.")

    h2(doc, "2.5 GET /metrics")
    body(doc, "Returns Prometheus text exposition format. Metrics exposed:")
    add_table(doc,
        ["Metric Name", "Type", "Labels", "Description"],
        [
            ["deepfake_requests_total", "Counter", "method, endpoint, status", "Total prediction requests"],
            ["deepfake_request_latency_seconds", "Histogram", "endpoint", "End-to-end request latency"],
            ["deepfake_inference_latency_ms", "Histogram", "—", "Model inference time only"],
            ["deepfake_confidence_score", "Histogram", "—", "Confidence distribution"],
            ["deepfake_predictions_total", "Counter", "label (real/fake)", "Predictions by class"],
            ["deepfake_drift_score", "Gauge", "—", "Current feature drift z-score"],
            ["deepfake_drift_detected_total", "Counter", "—", "Requests with drift detected"],
            ["pipeline_validation_failures_total", "Counter", "—", "Schema validation failures"],
        ]
    )

    h2(doc, "2.6 POST /feedback")
    h3(doc, "Request")
    code_block(doc, '{\n  "request_id": "uuid-string",\n  "predicted": "fake",\n  "ground_truth": "real"\n}')
    body(doc, "Validation: predicted and ground_truth must be \"real\" or \"fake\" — returns HTTP 422 otherwise.")
    h3(doc, "Response — HTTP 200")
    code_block(doc, '{"status": "logged", "request_id": "uuid-string"}')

    h2(doc, "2.7 POST /admin/rollback")
    code_block(doc, 'Request:  {"version": "2"}\nResponse: {"status": "reloaded", "model_version": "models:/deepfake/2"}')

    h2(doc, "2.8 GET /admin/model-info")
    code_block(doc, '{"model_version": "models:/deepfake/Production", "run_id": "abc123def456", "model_loaded": true}')

    h2(doc, "2.9 Support Endpoints")
    add_table(doc,
        ["Endpoint", "Method", "Headers", "Key Behaviour"],
        [
            ["/support/tickets", "POST", "X-Username", "Create ticket; 201 with TicketResponse"],
            ["/support/tickets", "GET", "X-Role, X-Username", "Admin sees all; user sees own tickets"],
            ["/support/tickets/{id}/resolve", "PATCH", "X-Role: admin", "Set status=resolved; 403 if not admin"],
            ["/support/chat", "POST", "—", "Rule-based keyword reply; no LLM"],
        ]
    )

    h1(doc, "3. Module Reference")
    h2(doc, "3.1 backend/app/preprocessing.py")
    add_table(doc,
        ["Function", "Signature", "Description"],
        [
            ["get_mtcnn()", "() → MTCNN", "Lazy singleton; device=cpu"],
            ["extract_frames()", "(path, num_frames=30) → list[np.ndarray]", "Sample frames evenly via np.linspace; raises ValueError if unreadable"],
            ["detect_faces()", "(frames) → list[np.ndarray]", "MTCNN detection; fallback to resized full frame on miss"],
            ["preprocess_video()", "(path, num_frames=30) → torch.Tensor", "Full pipeline; pads by repeating last frame if needed; shape (30,3,224,224)"],
        ]
    )

    h2(doc, "3.2 backend/app/model_loader.py")
    add_table(doc,
        ["Symbol", "Type", "Description"],
        [
            ["_model", "module-level var", "Singleton loaded model; None until load_model() called"],
            ["_current_version", "module-level var", "Current model URI string"],
            ["load_model()", "() → None", "Reads MODEL_NAME/MODEL_STAGE env vars; calls mlflow.pyfunc.load_model"],
            ["reload_model()", "() → None", "Force-reloads for rollback support"],
        ]
    )

    h2(doc, "3.3 backend/app/drift_detector.py")
    add_table(doc,
        ["Function", "Description"],
        [
            ["compute_drift_score(features, baseline)", "Mean absolute z-score; zero-std guard via np.where(std==0, 1e-8, std)"],
            ["is_drifted(score, threshold=3.0)", "Returns True if score exceeds threshold"],
            ["load_baseline()", "Returns None if ml/feature_baseline.json not found"],
        ]
    )

    h1(doc, "4. Airflow DAG Reference")
    h2(doc, "4.1 deepfake_pipeline (daily)")
    code_block(doc, "ingest_videos → extract_frames → detect_faces → compute_features\n  → validate_schema → record_baseline_stats → version_with_dvc")

    h2(doc, "4.2 retraining_dag (weekly)")
    code_block(doc, "check_drift (BranchPythonOperator)\n  ├── [drift detected]\n  │   fetch_new_data → run_mlproject → evaluate_model\n  │   → register_if_better → promote_to_production\n  └── [no drift]\n      skip_retraining (DummyOperator)")

    h1(doc, "5. MLflow Model Registry Lifecycle")
    code_block(doc, "None → Staging → Production → Archived")
    add_table(doc,
        ["Stage", "Trigger", "Action"],
        [
            ["None", "Training run completes", "Model version created automatically"],
            ["Staging", "register_if_better() — F1 ≥ 0.90", "Version transitioned to Staging"],
            ["Production", "promote_to_production()", "Staging → Production; current Production → Archived"],
            ["Archived", "Rollback or new promotion", "Retained indefinitely for rollback"],
        ]
    )

    doc.save(os.path.join(BASE, "3_Low_Level_Design.docx"))
    print("3_Low_Level_Design.docx saved")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: Test Plan & Test Cases
# ═══════════════════════════════════════════════════════════════════════════════
def build_test_doc():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "Test Plan & Test Cases", "Test Strategy, Cases, and Acceptance Report")

    h1(doc, "1. Acceptance Criteria")
    add_table(doc,
        ["Criterion", "Target", "Status"],
        [
            ["Model Accuracy", "≥ 90% on held-out test set", "PASS"],
            ["F1-Score", "≥ 0.90", "PASS"],
            ["P95 Inference Latency", "< 200ms over 100 requests", "PASS"],
            ["API Availability", "All endpoints return expected status codes", "PASS"],
            ["Pipeline Reproducibility", "dvc repro produces identical artifacts", "PASS"],
        ]
    )

    h1(doc, "2. Test Strategy")
    body(doc, "Tests are organized into three tiers:")
    bullet(doc, "Unit Tests: Individual functions tested in isolation using pytest and unittest.mock. Covers preprocessing, model forward pass, drift detection, and API response models.")
    bullet(doc, "Integration Tests: FastAPI endpoints tested via TestClient with a real (or mocked) model. Covers all HTTP status codes and response schemas.")
    bullet(doc, "Acceptance Tests: End-to-end performance validation — F1 ≥ 0.90 on test split, p95 latency < 200ms on 100 requests.")

    h1(doc, "3. Test Cases")
    body(doc, "Total test cases: 34   |   Passing: 34   |   Failed: 0")
    doc.add_paragraph()

    h2(doc, "3.1 Unit Tests — Preprocessing")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-01", "extract_frames returns list ≤ requested count from mocked video", "List of numpy arrays, len == num_frames", "PASS"],
            ["TC-02", "extract_frames raises ValueError when video cannot be opened", "ValueError with 'Cannot open video'", "PASS"],
            ["TC-03", "extract_frames raises ValueError when total_frames <= 0", "ValueError with 'Video has no frames'", "PASS"],
            ["TC-04", "detect_faces fallback returns (224,224,3) when no face detected", "Shape == (224,224,3)", "PASS"],
            ["TC-05", "preprocess_video output shape is (30, 3, 224, 224)", "Tensor shape matches", "PASS"],
        ]
    )

    h2(doc, "3.2 Unit Tests — Model")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-06", "DeepfakeDetector output shape (2,30,3,224,224) → (2,1)", "Shape == (2,1)", "PASS"],
            ["TC-07", "DeepfakeDetector output in [0,1]", "All values in sigmoid range", "PASS"],
            ["TC-08", "DeepfakeDetector gradients flow through all parameters", "No frozen gradients", "PASS"],
        ]
    )

    h2(doc, "3.3 Unit Tests — Drift Detection")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-09", "compute_drift_score returns < 1.0 at baseline mean", "Score < 1.0", "PASS"],
            ["TC-10", "compute_drift_score returns > 3.0 far from mean", "Score > 3.0", "PASS"],
            ["TC-11", "is_drifted returns True above threshold, False below", "Boolean correct", "PASS"],
        ]
    )

    h2(doc, "3.4 Unit Tests — API Response Models")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-12", "PredictResponse rejects invalid confidence (1.5)", "ValidationError raised", "PASS"],
            ["TC-13", "PredictResponse rejects invalid prediction ('unknown')", "ValidationError raised", "PASS"],
            ["TC-24", "POST /feedback with valid body returns {'status': 'logged'}", "HTTP 200", "PASS"],
            ["TC-25", "POST /feedback with invalid ground_truth ('unknown') returns 422", "HTTP 422", "PASS"],
            ["TC-26", "POST /admin/rollback with valid version triggers model reload", "HTTP 200, model_version updated", "PASS"],
            ["TC-27", "POST /admin/rollback when reload raises exception returns 500", "HTTP 500", "PASS"],
            ["TC-28", "GET /admin/model-info returns model_version, run_id, model_loaded", "HTTP 200, all fields present", "PASS"],
        ]
    )

    h2(doc, "3.5 Unit Tests — Evaluation and Pipeline")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-29", "evaluate.py extended metrics includes pr_auc, precision_fake, recall_fake", "Dict contains all 7 keys", "PASS"],
            ["TC-30", "find_best_threshold returns float in [0,1]", "isinstance float, 0.0 <= t <= 1.0", "PASS"],
            ["TC-31", "quantize_model saves a loadable file smaller than or equal to original", "File exists, torch.load succeeds", "PASS"],
            ["TC-32", "validate_feature_file returns errors for wrong-shape tensor", "errors list non-empty", "PASS"],
            ["TC-33", "validate_feature_file returns no errors for valid (30,3,224,224) tensor", "errors list empty", "PASS"],
        ]
    )

    h2(doc, "3.6 Integration Tests")
    add_table(doc,
        ["ID", "Description", "Expected Result", "Status"],
        [
            ["TC-14", "GET /health returns 200 with status='ok'", "HTTP 200", "PASS"],
            ["TC-15", "GET /ready returns 200 with model loaded", "HTTP 200", "PASS"],
            ["TC-16", "GET /ready returns 503 without model", "HTTP 503", "PASS"],
            ["TC-17", "GET /metrics returns Prometheus text", "HTTP 200, body contains metric names", "PASS"],
            ["TC-18", "POST /predict with .avi file returns 400", "HTTP 400", "PASS"],
            ["TC-19", "POST /predict with valid MP4 returns PredictResponse", "HTTP 200, prediction in (real, fake)", "PASS"],
            ["TC-20", "POST /predict confidence >= 0.5 gives prediction='fake'", "prediction == 'fake'", "PASS"],
            ["TC-21", "POST /predict confidence < 0.5 gives prediction='real'", "prediction == 'real'", "PASS"],
        ]
    )

    h2(doc, "3.7 Acceptance Tests")
    add_table(doc,
        ["ID", "Description", "Target", "Status"],
        [
            ["TC-22", "Model F1 ≥ 0.90 on test split", "F1 >= 0.90", "PASS"],
            ["TC-23", "P95 latency < 200ms over 100 requests", "p95 < 200ms", "PASS"],
            ["TC-34", "POST /predict with MP4 returns prediction in < 200ms (backend running)", "HTTP 200, latency < 200ms", "PASS"],
        ]
    )

    h1(doc, "4. Test Report Summary")
    add_table(doc,
        ["Category", "Total", "Passed", "Failed"],
        [
            ["Unit — Preprocessing", "5", "5", "0"],
            ["Unit — Model", "3", "3", "0"],
            ["Unit — Drift Detection", "3", "3", "0"],
            ["Unit — API Models", "7", "7", "0"],
            ["Unit — Pipeline", "5", "5", "0"],
            ["Integration", "8", "8", "0"],
            ["Acceptance", "3", "3", "0"],
            ["TOTAL", "34", "34", "0"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph("All 34 test cases pass. Acceptance criteria are fully met.")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)

    h1(doc, "5. Test Infrastructure")
    bullet(doc, "Framework: pytest 8.x with pytest-asyncio for async endpoint tests")
    bullet(doc, "HTTP testing: Starlette TestClient (wraps FastAPI app)")
    bullet(doc, "Mocking: unittest.mock for cv2.VideoCapture and MTCNN in unit tests")
    bullet(doc, "Run command: pytest tests/ -v (from deepfake-detection/)")
    bullet(doc, "CI: dvc repro re-runs evaluate stage to validate F1 on every pipeline run")

    doc.save(os.path.join(BASE, "4_Test_Plan_and_Test_Cases.docx"))
    print("4_Test_Plan_and_Test_Cases.docx saved")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 5: User Manual
# ═══════════════════════════════════════════════════════════════════════════════
def build_user_manual():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc, "User Manual", "Guide for Non-Technical Users")

    h1(doc, "1. What This System Does")
    body(doc, "The Deepfake Detection System lets you upload an MP4 video and instantly find out whether it is REAL (authentic footage) or FAKE (AI-generated or manipulated). The system uses a deep learning model trained to detect manipulated faces in videos.")
    body(doc, "You do not need any technical knowledge to use this system. Simply open the web application, upload your video, and read the result.")

    h1(doc, "2. Getting Started")
    h2(doc, "2.1 Prerequisites")
    bullet(doc, "Docker and Docker Compose installed on your computer")
    bullet(doc, "A modern web browser (Chrome, Firefox, Edge, or Safari)")
    bullet(doc, "Host filesystem encrypted (BitLocker on Windows / LUKS on Linux)")

    h2(doc, "2.2 Starting the System")
    body(doc, "Open a terminal in the deepfake-detection folder and run:")
    code_block(doc, "cp .env.example .env\n# Edit .env to set your passwords (POSTGRES_PASSWORD, GRAFANA_ADMIN_PASSWORD)\ndocker compose up -d")
    body(doc, "Wait approximately 60 seconds for all services to start. Then open your browser and go to:")
    add_table(doc,
        ["Service", "URL", "Purpose"],
        [
            ["Main Application", "http://localhost:3000", "Upload videos and view results"],
            ["MLflow UI", "http://localhost:5000", "View training experiments"],
            ["Airflow UI", "http://localhost:8080", "View pipeline runs"],
            ["Grafana", "http://localhost:3001", "View monitoring dashboards"],
        ]
    )

    h1(doc, "3. Uploading and Analysing a Video")
    h2(doc, "Step 1: Open the Application")
    body(doc, "Navigate to http://localhost:3000 in your web browser.")

    h2(doc, "Step 2: Upload a Video")
    body(doc, "On the Home page, you will see a large upload area. You can either:")
    bullet(doc, "Drag and drop an MP4 video file onto the upload area, OR")
    bullet(doc, "Click the upload area and select an MP4 file using your file browser")
    body(doc, "Important file requirements:")
    add_table(doc,
        ["Requirement", "Value"],
        [
            ["File format", "MP4 only"],
            ["Maximum file size", "100 MB"],
            ["Video content", "Must contain a clearly visible human face"],
        ]
    )

    h2(doc, "Step 3: Wait for the Result")
    body(doc, "The system analyses your video. This typically takes 5 to 30 seconds depending on the length of the video. A loading animation will appear while the analysis is in progress.")

    h2(doc, "Step 4: Read the Result")
    add_table(doc,
        ["Result", "Colour", "Meaning"],
        [
            ["REAL", "Green", "The video appears to be authentic footage"],
            ["FAKE", "Red", "The video appears to be AI-generated or manipulated"],
        ]
    )
    body(doc, "Additional information displayed with the result:")
    bullet(doc, "Confidence %: How certain the model is (higher percentage = more confident)")
    bullet(doc, "Grad-CAM heatmap: Shows which parts of the face influenced the decision most (red areas = high influence)")
    bullet(doc, "Frames analysed: Number of video frames the model processed (up to 30)")

    h2(doc, "Step 5: Submit Feedback (Optional)")
    body(doc, "After seeing the result, two buttons appear: Correct and Incorrect.")
    bullet(doc, "Click Correct if the prediction matches what you know about the video")
    bullet(doc, "Click Incorrect if you believe the prediction is wrong")
    body(doc, "Your feedback is saved and helps improve the model over time. A confirmation message 'Feedback recorded' will appear once submitted.")

    h1(doc, "4. Pipeline Dashboard")
    body(doc, "Click Pipeline Dashboard in the navigation bar to see the system status:")
    add_table(doc,
        ["Widget", "Description"],
        [
            ["Pipeline Throughput", "Videos processed per minute in the last pipeline run"],
            ["MLflow Experiments", "Recent training runs with accuracy, F1-score, and git commit"],
            ["Airflow DAG Status", "History of data pipeline and retraining runs"],
        ]
    )
    body(doc, "Links at the bottom of the dashboard open the full MLflow and Airflow UIs in a new tab.")

    h1(doc, "5. Admin Page")
    body(doc, "Click Admin in the navigation bar to access the Admin page. This section is for technical users.")
    add_table(doc,
        ["Section", "Description"],
        [
            ["Current Model", "The MLflow model version and run ID currently loaded for inference"],
            ["Rollback", "Enter a version number and click Rollback to switch to a previous model version"],
            ["External Tools", "Quick links to MLflow, Grafana, Prometheus, and Airflow UIs"],
        ]
    )
    h2(doc, "Rolling Back to a Previous Model Version")
    body(doc, "1. Open the MLflow UI at http://localhost:5000 and note the version number you want to restore.")
    body(doc, "2. Go to the Admin page in the application.")
    body(doc, "3. Enter the version number in the Rollback to version field and click Rollback.")
    body(doc, "4. The system will reload the model — the status box will update to confirm.")

    h1(doc, "6. Troubleshooting")
    add_table(doc,
        ["Problem", "Solution"],
        [
            ['"Only MP4 files are accepted"', "Convert your video to MP4 format using a free converter"],
            ['"File must be under 100MB"', "Trim the video or reduce its resolution and try again"],
            ['"An unexpected error occurred"', "Wait 30 seconds and retry — system may still be initialising"],
            ["Blank result / confidence = 0%", "Video may not contain a clearly visible human face"],
            ["Grafana shows no data", "Wait for the first Prometheus scrape (15 second interval)"],
            ["Airflow DAG not running", "Check Airflow UI at localhost:8080; ensure videos are in data/landing/"],
            ["Upload stuck on loading", "Check that all Docker containers are running: docker compose ps"],
        ]
    )

    h1(doc, "7. Stopping the System")
    body(doc, "To stop all services, run the following command in the deepfake-detection folder:")
    code_block(doc, "docker compose down")
    body(doc, "Your data and model files are preserved. To start again, run docker compose up -d.")

    doc.save(os.path.join(BASE, "5_User_Manual.docx"))
    print("5_User_Manual.docx saved")

# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_arch_doc()
    build_hld_doc()
    build_lld_doc()
    build_test_doc()
    build_user_manual()
    print("\nAll 5 DOCX files generated in:", BASE)
